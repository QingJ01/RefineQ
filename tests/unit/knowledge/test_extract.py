"""Tests for local study-material text extraction."""

from __future__ import annotations

from io import BytesIO

import fitz
import pytest
from docx import Document

from refineq.knowledge import extract as extraction
from refineq.knowledge.extract import MaterialExtractionError, extract_text


def test_extracts_utf8_text_and_markdown() -> None:
    assert (
        extract_text("notes.txt", "text/plain", "导数 derivative".encode())
        == "导数 derivative"
    )
    assert extract_text("notes.md", "text/markdown", b"# Limits\nContinuity") == (
        "# Limits\nContinuity"
    )


def test_extracts_pdf_text() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "The derivative measures local change.")
    payload = document.tobytes()
    document.close()

    extracted = extract_text("calculus.pdf", "application/pdf", payload)

    assert "derivative measures local change" in extracted


def test_extracts_docx_paragraphs() -> None:
    document = Document()
    document.add_heading("Calculus", level=1)
    document.add_paragraph("The chain rule composes derivatives.")
    buffer = BytesIO()
    document.save(buffer)

    extracted = extract_text(
        "calculus.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )

    assert extracted == "Calculus\nThe chain rule composes derivatives."


def test_pdf_page_budget_is_enforced_before_page_text_extraction() -> None:
    document = fitz.open()
    document.new_page()
    document.new_page()
    payload = document.tobytes()
    document.close()

    with pytest.raises(MaterialExtractionError, match="page"):
        extract_text(
            "large.pdf",
            "application/pdf",
            payload,
            limits=extraction.ExtractionLimits(max_pdf_pages=1),
        )


def test_docx_expanded_byte_and_compression_ratio_budgets_are_enforced() -> None:
    document = Document()
    document.add_paragraph("Repeated study note. " * 500)
    buffer = BytesIO()
    document.save(buffer)

    with pytest.raises(MaterialExtractionError, match="expanded"):
        extract_text(
            "large.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
            limits=extraction.ExtractionLimits(max_docx_expanded_bytes=100),
        )
    with pytest.raises(MaterialExtractionError, match="compression"):
        extract_text(
            "large.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
            limits=extraction.ExtractionLimits(max_docx_compression_ratio=1.0),
        )


def test_extracted_character_budget_applies_to_plain_text() -> None:
    with pytest.raises(MaterialExtractionError, match="character"):
        extract_text(
            "large.txt",
            "text/plain",
            b"01234567890",
            limits=extraction.ExtractionLimits(max_extracted_chars=10),
        )
